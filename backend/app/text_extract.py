"""
Extract raw text from an uploaded resume file (PDF, DOCX, or TXT).
"""

from __future__ import annotations

import io

import docx
import pdfplumber


class UnsupportedFileType(Exception):
    pass


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_text_from_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    if lower.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    raise UnsupportedFileType(
        f"Unsupported file type for '{filename}'. Please upload a PDF, DOCX, or TXT resume."
    )
